import os
import sys
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_diagrams():
    print("Generating diagrams...")
    # 1. System Architecture Diagram
    fig, ax = plt.subplots(figsize=(8, 3.2))
    ax.axis('off')

    # Box details: text, position (x, y), width, height, bg_color, border_color
    boxes = [
        ("Client Browser\n(Explore/Compare/Dashboard)", 0.03, 0.4, 0.22, 0.22, '#E0F2FE', '#0284C7'),
        ("Next.js Frontend\n(React Router + TS)", 0.32, 0.4, 0.22, 0.22, '#F0FDF4', '#16A34A'),
        ("ASP.NET Core Web API\n(C# Backend Engine)", 0.61, 0.4, 0.22, 0.22, '#FAF5FF', '#7E22CE'),
        ("PostgreSQL DB\n(mcc_portfolio)", 0.88, 0.4, 0.10, 0.22, '#FEF2F2', '#DC2626')
    ]

    # Draw boxes
    for text, x, y, w, h, bg, border in boxes:
        rect = plt.Rectangle((x, y), w, h, facecolor=bg, edgecolor=border, lw=2, transform=ax.transAxes, zorder=3)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2, text, color='#0F172A', fontsize=8, fontweight='bold',
                ha='center', va='center', wrap=True, transform=ax.transAxes, zorder=4)

    # Draw arrows
    arrows = [
        (0.25, 0.51, 0.32, 0.51),
        (0.54, 0.51, 0.61, 0.51),
        (0.83, 0.51, 0.88, 0.51)
    ]
    for x1, y1, x2, y2 in arrows:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="<->", color='#475569', lw=2, zorder=2),
                    xycoords='axes fraction', textcoords='axes fraction')

    # Draw local AI Engine note attached to Web API
    rect_ai = plt.Rectangle((0.63, 0.05), 0.18, 0.18, facecolor='#FEF3C7', edgecolor='#D97706', lw=1.5, ls='--', transform=ax.transAxes, zorder=3)
    ax.add_patch(rect_ai)
    ax.text(0.72, 0.14, "Rule Engine\n(AI Assistant)", color='#0F172A', fontsize=7, fontweight='bold',
            ha='center', va='center', wrap=True, transform=ax.transAxes, zorder=4)
    # Arrow to API
    ax.annotate("", xy=(0.72, 0.4), xytext=(0.72, 0.23),
                arrowprops=dict(arrowstyle="<->", color='#64748B', lw=1.5, ls='--', zorder=2),
                xycoords='axes fraction', textcoords='axes fraction')

    plt.title("MCC Portfolio System - Multi-Tier Architecture", fontsize=10, fontweight='bold', pad=10, color='#1E293B')
    plt.tight_layout()
    plt.savefig('system_architecture.png', dpi=300)
    plt.close()

    # 2. Database Relationships Diagram
    fig, ax = plt.subplots(figsize=(9.5, 5))
    ax.axis('off')

    # Table specifications
    tables = [
        ("Users\n- Id [PK]\n- Username\n- Email\n- Role", 0.03, 0.72, 0.20, 0.22, '#F1F5F9', '#475569'),
        ("StudentProfiles\n- Id [PK]\n- UserId [FK]\n- FullName\n- Bio\n- Department\n- Theme\n- Approved", 0.36, 0.45, 0.25, 0.38, '#ECFDF5', '#059669'),
        ("Projects\n- Id [PK]\n- ProfileId [FK]\n- Title\n- TechStack\n- GithubUrl", 0.75, 0.81, 0.21, 0.17, '#EFF6FF', '#2563EB'),
        ("Certifications\n- Id [PK]\n- ProfileId [FK]\n- Name\n- Issuer\n- CredentialUrl", 0.75, 0.61, 0.21, 0.17, '#EFF6FF', '#2563EB'),
        ("ResearchPapers\n- Id [PK]\n- ProfileId [FK]\n- Title\n- JournalName\n- PublishDate", 0.75, 0.41, 0.21, 0.17, '#EFF6FF', '#2563EB'),
        ("Hackathons / Extra\n- Id [PK]\n- ProfileId [FK]\n- EventName\n- Prize", 0.75, 0.21, 0.21, 0.17, '#EFF6FF', '#2563EB'),
        ("ThemeConfigs\n- Id [PK]\n- Name\n- PrimaryColor\n- IsEnabled", 0.03, 0.35, 0.20, 0.20, '#FFFBEB', '#D97706'),
        ("Institutions\n- Id [PK]\n- Name\n- Address\n- WebsiteUrl", 0.03, 0.08, 0.20, 0.20, '#FFFBEB', '#D97706')
    ]

    for text, x, y, w, h, bg, border in tables:
        rect = plt.Rectangle((x, y), w, h, facecolor=bg, edgecolor=border, lw=1.5, transform=ax.transAxes, zorder=3)
        ax.add_patch(rect)
        ax.text(x + 0.02, y + h - 0.03, text, color='#0F172A', fontsize=7.5, fontname='monospace',
                ha='left', va='top', wrap=True, transform=ax.transAxes, zorder=4)

    # Drawing relationships (arrows)
    # Users -> StudentProfiles (1:1)
    ax.annotate("", xy=(0.36, 0.75), xytext=(0.23, 0.82),
                arrowprops=dict(arrowstyle="->", connectionstyle="arc3,rad=-0.1", color='#64748B', lw=1.5, zorder=2),
                xycoords='axes fraction', textcoords='axes fraction')

    # StudentProfiles -> Projects (1:N)
    ax.annotate("", xy=(0.75, 0.88), xytext=(0.61, 0.78),
                arrowprops=dict(arrowstyle="->", connectionstyle="arc3,rad=0.1", color='#64748B', lw=1.5, zorder=2),
                xycoords='axes fraction', textcoords='axes fraction')

    # StudentProfiles -> Certs (1:N)
    ax.annotate("", xy=(0.75, 0.70), xytext=(0.61, 0.65),
                arrowprops=dict(arrowstyle="->", color='#64748B', lw=1.5, zorder=2),
                xycoords='axes fraction', textcoords='axes fraction')

    # StudentProfiles -> Research (1:N)
    ax.annotate("", xy=(0.75, 0.50), xytext=(0.61, 0.58),
                arrowprops=dict(arrowstyle="->", color='#64748B', lw=1.5, zorder=2),
                xycoords='axes fraction', textcoords='axes fraction')

    # StudentProfiles -> Hackathons (1:N)
    ax.annotate("", xy=(0.75, 0.30), xytext=(0.61, 0.50),
                arrowprops=dict(arrowstyle="->", connectionstyle="arc3,rad=-0.1", color='#64748B', lw=1.5, zorder=2),
                xycoords='axes fraction', textcoords='axes fraction')

    plt.title("MCC Portfolio PostgreSQL Database Relationships Schema", fontsize=10, fontweight='bold', pad=10, color='#1E293B')
    plt.tight_layout()
    plt.savefig('database_relations.png', dpi=300)
    plt.close()
    print("Diagrams generated successfully!")

def build_docx():
    print("Building Docx file...")
    doc = Document()

    # Style Configurations
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate-700

    # Custom color palette constants (MCC Maroon)
    MCC_COLOR = RGBColor(0x80, 0x00, 0x20) # Maroon
    ACCENT_COLOR = RGBColor(0x02, 0x84, 0x0C7) # Slate Blue

    # Helper function to add headers
    def add_custom_heading(text, level, space_before=Pt(12), space_after=Pt(6)):
        h = doc.add_heading(text, level=level)
        h.paragraph_format.space_before = space_before
        h.paragraph_format.space_after = space_after
        h.paragraph_format.keep_with_next = True
        
        run = h.runs[0]
        run.font.name = 'Segoe UI'
        run.font.bold = True
        
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = MCC_COLOR
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = ACCENT_COLOR
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        return h

    # 1. Header Cover Info
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(40)
    title_p.paragraph_format.space_after = Pt(4)
    run_inst = title_p.add_run("MADRAS CHRISTIAN COLLEGE\n")
    run_inst.font.name = 'Segoe UI'
    run_inst.font.size = Pt(12)
    run_inst.font.bold = True
    run_inst.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    run_main = title_p.add_run("STUDENT PORTFOLIO PLATFORM\n")
    run_main.font.name = 'Segoe UI'
    run_main.font.size = Pt(24)
    run_main.font.bold = True
    run_main.font.color.rgb = MCC_COLOR

    run_sub = title_p.add_run("Project Implementation & Departmental Registry Report")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # Metadata Table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    
    meta_data = [
        ("Prepared For:", "Department of Computer Science / Academic Audit Committee"),
        ("Date of Report:", "June 9, 2026"),
        ("Software Stack:", "ASP.NET Core Web API, PostgreSQL, Next.js (TypeScript)"),
        ("Deployment Status:", "Ready for Production (0 errors, dev servers online)")
    ]

    for i, (label, val) in enumerate(meta_data):
        row = table.rows[i]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
        row.cells[0].width = Inches(2.0)
        
        row.cells[1].paragraphs[0].add_run(val)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)
        row.cells[1].width = Inches(4.5)

    doc.add_page_break()

    # SECTION 1
    add_custom_heading("1. Executive Summary", level=1)
    p = doc.add_paragraph(
        "The Madras Christian College (MCC) Student Portfolio Platform is an advanced web-based registry and "
        "digital showcase application designed to track and publish student achievements. The platform addresses "
        "key institutional documentation challenges, enabling students to log their projects, certifications, research "
        "publications, hackathons, and community services. In doing so, it serves as a central registry that aligns "
        "directly with the rigorous document auditing parameters established by NAAC and NIRF accreditation teams."
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(10)

    p2 = doc.add_paragraph(
        "By hosting a dynamic gallery with live profile sharing, side-by-side competency comparisons, and professional "
        "PDF resume exporting, the platform bridges the gap between academic progress and industry recruitment, "
        "advancing the digital footprint of the Computer Science Department."
    )
    p2.paragraph_format.line_spacing = 1.15
    p2.paragraph_format.space_after = Pt(10)

    # SECTION 2
    add_custom_heading("2. System Architecture & Tech Stack", level=1)
    doc.add_paragraph(
        "The platform utilizes a modern, decoupled multi-tier architecture to deliver high-performance user interfaces "
        "and securely managed relational database transactions. Standard JSON web tokens (JWT) handle student "
        "and administrator authorization states, while Entity Framework Core handles communication with a local PostgreSQL server."
    )

    # Add system architecture image
    if os.path.exists("system_architecture.png"):
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.add_run().add_picture("system_architecture.png", width=Inches(6.0))
        caption = img_p.add_run("\nFigure 2.1: Multi-Tier System Architecture & Routing Flow")
        caption.font.size = Pt(9)
        caption.font.italic = True
        caption.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        img_p.paragraph_format.space_after = Pt(12)

    # Technical specifics bullet list
    doc.add_paragraph("Core Technology Specifications:").bold = True
    bullets = [
        ("Frontend Application", "Next.js App Router (React, TypeScript) styling via CSS theme configurations and layout modules."),
        ("Backend Web API", "ASP.NET Core (C#) utilizing Controllers, custom DTO bindings, Cors middleware configuration, and Jwt security pipeline."),
        ("Database Layer", "PostgreSQL running locally at port 5432, storing tables for profiles, records, notifications, and institution themes."),
        ("AI Analytics Engine", "A C# logic rules engine assessing tech stacks, portfolio completeness scores, and crafting dynamic letters.")
    ]
    for title, desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        run_title = bp.add_run(title + ": ")
        run_title.bold = True
        run_title.font.color.rgb = ACCENT_COLOR
        bp.add_run(desc)

    doc.add_page_break()

    # SECTION 3
    add_custom_heading("3. Database Schema & Models", level=1)
    doc.add_paragraph(
        "The relational database schema is structured to accommodate the diverse range of milestones "
        "encountered by undergraduate and postgraduate students. Cascading deletes are enforced on the User profile level "
        "to ensure student data integrity, matching standard privacy and compliance rules."
    )

    # Add database relations image
    if os.path.exists("database_relations.png"):
        img_p2 = doc.add_paragraph()
        img_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p2.add_run().add_picture("database_relations.png", width=Inches(6.2))
        caption2 = img_p2.add_run("\nFigure 3.1: PostgreSQL Relational Database Schema & Primary/Foreign Key Mappings")
        caption2.font.size = Pt(9)
        caption2.font.italic = True
        caption2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        img_p2.paragraph_format.space_after = Pt(12)

    # SECTION 4
    add_custom_heading("4. Key Platform Features & Modules", level=1)
    
    add_custom_heading("4.1 Interactive Accomplishment Milestones", level=2)
    doc.add_paragraph(
        "To present a high-fidelity visual summary of students' accomplishments, the public portfolio route incorporates:\n"
        "1. Competency Skill Radar: An SVG radar polygon parsing coding project tech stacks, certifications, and creative "
        "contributions into five domain scores: Frontend, Backend, Databases, DevOps & Tools, and Creative/Research.\n"
        "2. Chronological Milestones Timeline: A vertical timeline sorting and presenting all student achievements. Supports "
        "filters (Projects, Credentials, Extracurriculars) and detail modal drawers."
    )

    add_custom_heading("4.2 AI Career Guidance Console", level=2)
    doc.add_paragraph(
        "Students are equipped with three helper functions driven by a local algorithmic rules engine:\n"
        "• Resume Completeness Critic: Automatically scores a profile from 0-100 based on modules completed and lists actionable suggestions.\n"
        "• Statement of Purpose (SOP) Generator: Automatically builds custom, formatted letters mapping projects to tone preferences.\n"
        "• Career Path Advisor: Identifies technical stack profiles and suggests roles, skills gaps, target universities, and scholarships."
    )

    add_custom_heading("4.3 Public Directory & Side-by-Side Comparison Engine", level=2)
    doc.add_paragraph(
        "For external visitors, admissions, or placement officers, the system provides:\n"
        "• Public Registry (/explore): Searchable directory allowing users to look up students by skills, name, or department.\n"
        "• Side-by-Side Comparison (/compare): Displays two students side-by-side. Highlights skill overlaps using a Dual-Overlay "
        "SVG Skill Radar Chart displaying both competency polygons in translucent Cyan and Purple layers."
    )

    add_custom_heading("4.4 Print-to-PDF Stylesheets & Resume Exporting", level=2)
    doc.add_paragraph(
        "A customized print layout configuration triggers on window.print() (Export Resume), instantly refactoring the "
        "profile structure into a clean, print-friendly 2-column professional resume layout. It automatically hides screen-only elements, "
        "inverts high-saturation background fills, and aligns contact parameters for high-fidelity printing."
    )

    add_custom_heading("4.5 Administrator Management Panel", level=2)
    doc.add_paragraph(
        "Department Heads and administrative users can check overall database counts, approve or revoke public profile listings, "
        "manage campus announcements, and download a single CSV backup containing all verified student details for record archiving."
    )

    # SECTION 5
    add_custom_heading("5. Accreditation & Academic Implementation Value", level=1)
    doc.add_paragraph(
        "The implementation of the MCC Student Portfolio Platform delivers measurable benefits for departments:\n"
        "1. NAAC/NIRF Criteria Data Audits: The administrative CSV exporter instantly maps student publications, projects, "
        "and community service logs, resolving manual information gathering.\n"
        "2. Streamlined Campus Recruitment: Provides recruiters with immediate, authenticated visibility into the students' coding portfolios, "
        "skills radar distribution charts, and project repositories.\n"
        "3. Departmental Showcasing: Organizes and preserves institutional memory, cataloging historical achievements, hackathons, and creative vector sketches."
    )

    # Save document
    filename = "MCC_Student_Portfolio_Platform_Project_Report.docx"
    doc.save(filename)
    print(f"Report saved to {filename}")

if __name__ == "__main__":
    create_diagrams()
    build_docx()
    # Clean up diagram images
    if os.path.exists("system_architecture.png"):
        os.remove("system_architecture.png")
    if os.path.exists("database_relations.png"):
        os.remove("database_relations.png")
    print("Done!")
